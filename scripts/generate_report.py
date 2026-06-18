"""Generate project technical report as Word document from JSON content."""

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
CONTENT_PATH = ROOT / "scripts" / "report_content.json"


def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def build_report():
    data = json.loads(CONTENT_PATH.read_text(encoding="utf-8"))
    doc = Document()

    # Style setup
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    for lv in range(1, 4):
        doc.styles[f"Heading {lv}"].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    # ── Cover ──
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(data["cover"]["title"] + "\n" + data["cover"]["subtitle"])
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(data["cover"]["tagline"])
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run(data["cover"]["meta"]).font.size = Pt(9)

    doc.add_page_break()

    # ── Helper functions ──
    def h(text, level=1):
        doc.add_heading(text, level=level)

    def p(text):
        doc.add_paragraph(text)

    def bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    def code(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def tbl(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Shading Accent 1")
        for i, hdr in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = hdr
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph()

    # ── 1. Overview ──
    h("一、项目概述")
    h("1.1 项目背景", 2)
    p(data["overview"]["background"])
    h("1.2 核心能力", 2)
    for c in data["overview"]["capabilities"]:
        bullet(c)

    # ── 2. Architecture ──
    h("二、项目架构")
    h("2.1 技术栈分层", 2)
    arch_headers = ["层级", "技术选型", "说明"]
    arch_rows = [[a["name"], a["tech"], a["desc"]] for a in data["architecture"]["layers"]]
    tbl(arch_headers, arch_rows)

    h("2.2 核心设计模式", 2)
    for pat in data["patterns"]:
        p(pat["name"])
        p(pat["desc"])

    # ── 3. Build Steps ──
    h("三、详细搭建流程")
    for step in data["steps"]:
        h(step["title"], 2)
        p(step["content"])

    # ── 4. Data Flow ──
    h("四、关键数据流向")
    h("4.1 完整对话请求链路", 2)
    for line in data["dataflow"]["flow"]:
        code(line)

    h("4.2 Tool Calling 状态机", 2)
    for state in data["state_machine"]["states"]:
        bullet(state)

    # ── 5. Database ──
    h("五、数据库设计")
    h("5.1 数据概览", 2)
    db_headers = ["表名", "行数", "核心字段"]
    tbl(db_headers, data["database"]["tables"])
    h("5.2 索引策略", 2)
    p(data["database"]["index_strategy"])

    # ── 6. Evaluation ──
    h("六、评测体系")
    h("6.1 评测结果", 2)
    eval_headers = ["Agent", "通过率", "工具准确率", "关键词覆盖", "平均响应"]
    tbl(eval_headers, data["eval_results"]["summary"])
    h("6.2 分析", 2)
    p(data["eval_results"]["insight"])

    # ── 7. Tests ──
    h("七、测试覆盖")
    tbl(["测试类", "用例", "状态"], data["tests"])

    # ── 8. Project Files ──
    h("八、项目文件清单")
    file_headers = ["文件", "说明"]
    file_rows = [[f[0], f[1]] for f in data["files"]]
    tbl(file_headers, file_rows)

    # Count actual line numbers
    for i, (fpath, desc) in enumerate(data["files"]):
        full = ROOT / fpath
        if full.exists():
            lines = len(full.read_text(encoding="utf-8").splitlines())
            file_rows[i][0] = f"{fpath} ({lines} lines)"

    # Rebuild table with line counts
    doc.add_paragraph()
    tbl(["文件（含行数）", "说明"], file_rows)

    total = sum(len((ROOT / f[0].split(" (")[0]).read_text(encoding="utf-8").splitlines())
                 for f in file_rows if (ROOT / f[0].split(" (")[0]).exists())
    p(f"总计 {len(file_rows)} 个核心文件，约 {total} 行 Python 代码。")

    # ── 9. Startup ──
    h("九、启动方式")
    for cmd in data["startup"]:
        if cmd.startswith("#"):
            p(cmd)
        else:
            code(cmd)

    # ── 10. Resume Highlights ──
    h("十、简历亮点提炼")
    for item in data["resume_highlights"]:
        bullet(item)

    # ── Save ──
    out = ROOT / "招投标智能体平台_技术报告.docx"
    doc.save(str(out))
    print(f"Report saved: {out}")
    print(f"Size: {os.path.getsize(out) / 1024:.0f} KB")


if __name__ == "__main__":
    build_report()
