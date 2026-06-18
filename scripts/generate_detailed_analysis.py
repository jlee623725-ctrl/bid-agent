"""Generate comprehensive project analysis as Word document."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = ROOT / "招投标智能体平台_详细代码分析.docx"


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_code_block(doc, text, font_size=Pt(8)):
    """Add a code block with monospace font and grey background."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = font_size
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    para = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return para


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    # Header
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = hdr
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacing after table
    return table


def build_document():
    doc = Document()

    # ── Style setup ──
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for lv in range(1, 5):
        heading_style = doc.styles[f"Heading {lv}"]
        heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        heading_style.font.name = "微软雅黑"
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ═══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("招投标智能体平台\n(Bid-Agent)")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("完整代码架构分析与文件详解")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = m.add_run("Multi-Agent + RAG + Tool Calling 智能体系统\n基于 DeepSeek API · FastAPI · Streamlit · SQLite · TF-IDF")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (placeholder)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("目录", level=1)
    toc_items = [
        "一、项目总览与宏观架构",
        "二、目录结构与文件组织",
        "三、核心引擎层 (agent/) 深度剖析",
        "  3.1 core.py — BidAgent 工具调用引擎",
        "  3.2 factory.py — Agent 注册工厂",
        "  3.3 supervisor.py — 备用 Supervisor 构造器",
        "  3.4 tools_bidding.py — 标讯分析工具集",
        "  3.5 tools_company.py — 企业画像工具集",
        "  3.6 tools_legal.py — 法规检索工具集",
        "  3.7 vector_store.py — TF-IDF 向量存储器",
        "四、展示与 API 层",
        "  4.1 api.py — FastAPI 后端服务",
        "  4.2 app.py — Streamlit 前端界面",
        "五、脚本工具层 (scripts/)",
        "  5.1 init_db.py — 数据库初始化管道",
        "  5.2 build_index.py — 向量索引构建",
        "  5.3 evaluate.py — 自动化评测框架",
        "  5.4 generate_report.py — Word 报告生成器",
        "六、测试与评测体系 (tests/)",
        "  6.1 test_agent.py — Agent 单元测试",
        "  6.2 test_tools.py — 工具函数集成测试",
        "  6.3 eval_data.json — 标注评测数据集",
        "七、数据层 (data/)",
        "八、配置文件与环境变量",
        "九、完整请求链路剖析",
        "十、设计模式与架构决策",
        "十一、关键技术细节与源码级分析",
        "十二、评测结果与性能分析",
        "十三、启动与运行指南",
        "十四、项目价值与技能亮点总结",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 1: 项目总览
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("一、项目总览与宏观架构", level=1)

    doc.add_paragraph(
        "招投标智能体平台 (Bid-Agent) 是一个基于大语言模型 (LLM) 的 Multi-Agent 智能分析系统，"
        "专为招投标领域设计。项目的核心目标是：让用户通过自然语言对话的方式，完成招标公告搜索、"
        "企业画像分析、竞争对手查找、法规政策检索等专业任务。"
    )

    doc.add_heading("1.1 业务背景", level=2)
    doc.add_paragraph(
        "在招投标行业，从业人员需要频繁进行以下操作：(1) 在海量招标公告中搜索相关标讯；"
        "(2) 分析特定企业的中标记录和竞争格局；(3) 查证法律法规的具体条款。"
        "这些操作通常需要跨多个系统、使用专业查询语法，门槛较高。"
        "本平台通过 AI Agent 技术，将这些专业操作封装为自然语言对话，大幅降低使用门槛。"
    )

    doc.add_heading("1.2 技术栈全景", level=2)
    add_table(doc,
        ["层级", "技术选型", "版本", "角色"],
        [
            ["LLM 服务", "DeepSeek API (OpenAI 兼容)", "deepseek-chat", "提供推理与工具调用能力"],
            ["Agent 引擎", "纯 Python 手写 Tool Calling 循环", "—", "管理对话状态与工具调度"],
            ["Web 前端", "Streamlit", ">=1.32.0", "多 Agent 对话界面"],
            ["API 后端", "FastAPI + Pydantic + Uvicorn", ">=0.100", "RESTful API 与会话管理"],
            ["RAG 检索引擎", "scikit-learn TF-IDF + Cosine Similarity", ">=1.0.0", "中文法律法规语义搜索"],
            ["数据库", "SQLite3 (row_factory=Row)", "内置", "结构化数据存储与查询"],
            ["数据处理", "Pandas", ">=2.0.0", "CSV 导入与数据清洗"],
            ["测试框架", "pytest + unittest.mock", ">=8.0.0", "单元测试与集成测试"],
            ["文档生成", "python-docx", "最新", "Word 技术报告生成"],
            ["环境管理", "python-dotenv", ">=1.0.0", ".env 环境变量加载"],
        ]
    )

    doc.add_heading("1.3 核心架构模式", level=2)
    doc.add_paragraph(
        "项目采用 Supervisor + Specialists 的 Multi-Agent 编排模式："
    )
    add_bullet(doc, "Supervisor Agent (智能助理)：拥有全部 9 个工具的访问权限，能够跨标讯、企业、法规三个领域进行编排推理。理解用户意图后，智能决策调用哪些工具并整合结果。")
    add_bullet(doc, "Bidding Analyst (标讯分析师)：专注招标公告搜索、中标趋势统计、公告详情查询，配备 3 个专用工具。")
    add_bullet(doc, "Company Profiler (企业画像师)：专注企业筛选、全景画像生成、竞争对手分析，配备 3 个专用工具。")
    add_bullet(doc, "Legal Advisor (法规顾问)：专注法律法规语义搜索、关键词检索、法条原文精确查询，配备 3 个专用工具。")

    doc.add_paragraph(
        "这种设计带来关键优势：Supervisor 可以跨领域整合信息（如'合肥建筑行业有哪些中标企业'"
        "需要同时调用企业搜索和标讯搜索），而 Specialist 在单一领域内更聚焦、成本更低。"
        "评测数据表明 Supervisor 的通过率 (86%) 比单个 Specialist (32%-42%) 高出 2 倍以上。"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 2: 目录结构
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("二、目录结构与文件组织", level=1)

    doc.add_paragraph(
        "项目遵循清晰的分层目录结构，核心业务逻辑在 agent/ 包中，数据在 data/ 目录，"
        "脚本工具在 scripts/ 目录，测试在 tests/ 目录，入口文件在项目根目录。"
    )

    add_table(doc,
        ["路径", "类型", "作用说明"],
        [
            [".env / .env.example", "配置文件", "环境变量：API密钥、数据库路径、日志级别"],
            ["requirements.txt", "依赖声明", "Python 包依赖清单 (8个核心包)"],
            ["CLAUDE.md", "项目规范", "AI 辅助开发的编码规范与约定"],
            ["api.py", "FastAPI 后端入口", "RESTful API 服务，5个端点，会话管理"],
            ["app.py", "Streamlit 前端入口", "多 Agent 对话 Web 界面，侧边栏切换"],
            ["agent/__init__.py", "包初始化", "空文件，标记 Python 包"],
            ["agent/core.py", "核心引擎", "BidAgent 类：Tool Calling 循环，212行"],
            ["agent/factory.py", "Agent 工厂", "AgentRegistry + 4个预注册 Agent，159行"],
            ["agent/supervisor.py", "备用构造器", "独立 Supervisor 创建函数，88行"],
            ["agent/tools_bidding.py", "标讯工具", "3个工具函数 + JSON Schema，195行"],
            ["agent/tools_company.py", "企业工具", "3个工具函数 + JSON Schema，267行"],
            ["agent/tools_legal.py", "法规工具", "3个工具函数 + JSON Schema，191行"],
            ["agent/vector_store.py", "向量存储", "TF-IDF 检索引擎，174行"],
            ["scripts/init_db.py", "构建脚本", "CSV → SQLite 自动化管道，315行"],
            ["scripts/build_index.py", "构建脚本", "TF-IDF 索引构建与验证，29行"],
            ["scripts/evaluate.py", "评测脚本", "50条用例的自动化评测框架，229行"],
            ["scripts/generate_report.py", "报告生成", "Word 技术报告生成器，184行"],
            ["scripts/report_content.json", "数据文件", "技术报告的结构化内容数据"],
            ["tests/conftest.py", "测试配置", "pytest 路径配置，5行"],
            ["tests/test_agent.py", "单元测试", "BidAgent 核心测试，8个用例，334行"],
            ["tests/test_tools.py", "集成测试", "工具函数测试，16个用例，124行"],
            ["tests/eval_data.json", "评测数据", "50条标注 QA 评测用例"],
            ["tests/eval_result_*.json", "评测结果", "4个 Agent 的评测结果详情"],
            ["data/bid_agent.db", "主数据库", "SQLite 数据库，约11张表 ~10000条记录"],
            ["data/*.csv", "原始数据", "11个 CSV 数据文件"],
            ["data/vector_index/", "索引目录", "TF-IDF 向量化器 + 矩阵 + 元数据"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 3: 核心引擎层
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("三、核心引擎层 (agent/) 深度剖析", level=1)

    # ── 3.1 core.py ──
    doc.add_heading("3.1 core.py — BidAgent 工具调用引擎 (212行)", level=2)

    doc.add_paragraph(
        "这是整个系统的核心引擎，实现了一个基于 OpenAI 兼容 API 的 Tool Calling 循环。"
        "该文件不依赖任何业务逻辑，是一个纯粹的、可复用的 Agent 框架。"
    )

    doc.add_heading("3.1.1 类结构与常量", level=3)
    add_table(doc,
        ["元素", "值/类型", "说明"],
        [
            ["MAX_ROUNDS", "10", "最大工具调用轮数，防止无限循环"],
            ["MAX_RETRIES", "2", "API 调用失败时的最大重试次数"],
            ["RETRY_BASE_DELAY", "1.0s", "指数退避重试的基础延迟"],
            ["ToolSchema", "List[Dict[str, Any]]", "OpenAI Function Calling JSON Schema 格式"],
            ["ToolRegistry", "Dict[str, Callable]", "工具名 → 处理函数的映射表"],
            ["BidAgent.client", "OpenAI 实例", "通过 api_key + base_url 初始化的 OpenAI 客户端"],
        ]
    )

    doc.add_heading("3.1.2 构造函数 __init__", level=3)
    doc.add_paragraph(
        "接收 4 个参数: system_prompt (系统提示词), tools (工具 JSON Schema 列表), "
        "tool_handlers (工具执行函数字典), model (模型名称，默认 deepseek-chat)。"
    )
    add_bullet(doc, "输入验证：检查 system_prompt 非空、tools 为列表类型")
    add_bullet(doc, "API 配置：从环境变量 DEEPSEEK_BASE_URL 和 DEEPSEEK_API_KEY 读取配置")
    add_bullet(doc, "客户端初始化：使用 openai.OpenAI(api_key, base_url) 创建客户端，兼容 DeepSeek API")
    add_bullet(doc, "日志记录：初始化时记录 model、base_url、tools 数量、handlers 数量")

    doc.add_heading("3.1.3 run() — 同步工具调用循环", level=3)
    doc.add_paragraph("这是核心执行方法，实现完整的 Tool Calling 循环。流程如下：")
    add_bullet(doc, "Step 1: 构建初始 messages = [system_prompt, user_input]")
    add_bullet(doc, "Step 2: 进入 for 循环 (最多 MAX_ROUNDS=10 轮)")
    add_bullet(doc, "Step 3: 每轮调用 _call_api(messages, tools)，记录耗时")
    add_bullet(doc, "Step 4: 检查 response.choices[0].finish_reason：")
    add_bullet(doc, "  - \"stop\": LLM 完成推理，返回 message.content 文本", level=1)
    add_bullet(doc, "  - \"tool_calls\": LLM 请求调用工具，执行 _execute_tool()，将工具结果追加到 messages，继续下一轮", level=1)
    add_bullet(doc, "  - 其他 (\"length\", \"content_filter\" 等): 返回已有内容", level=1)
    add_bullet(doc, "Step 5: 超轮数限制 → 返回 \"处理超时\"")

    doc.add_heading("3.1.4 run_stream() — 流式变体", level=3)
    doc.add_paragraph(
        "与 run() 逻辑相同但使用生成器模式。当 LLM 返回 finish_reason=\"stop\" 时，"
        "yield 内容文本。当模型请求 tool_calls 时，内部静默执行工具调用并继续循环，"
        "最终只流式输出文本结果。注意：DeepSeek API 在 tool_calls 时通常不使用 streaming。"
    )

    doc.add_heading("3.1.5 _call_api() — 带重试的 API 调用", level=3)
    doc.add_paragraph(
        "使用指数退避 (Exponential Backoff) 策略进行错误重试："
    )
    add_bullet(doc, "最多尝试 MAX_RETRIES+1 = 3 次 (初次+2次重试)")
    add_bullet(doc, "每次失败后延迟: RETRY_BASE_DELAY * 2^attempt (1s, 2s, 4s)")
    add_bullet(doc, "3次全部失败后抛出 RuntimeError")
    add_bullet(doc, "调用参数: model, messages, tools, stream=False")

    doc.add_heading("3.1.6 _execute_tool() — 工具执行器", level=3)
    doc.add_paragraph(
        "负责执行已注册的工具函数，包含完善的容错机制："
    )
    add_bullet(doc, "工具名未注册: 返回 \"Unknown tool: {name}\"，不崩溃，让 LLM 自修正")
    add_bullet(doc, "JSON 参数解析: json.loads(arguments)，解析失败则使用空字典")
    add_bullet(doc, "函数调用: handler(**args) (支持关键字参数展开)")
    add_bullet(doc, "异常捕获: 工具执行异常被捕获为文本返回，让 LLM 感知错误并换策略")
    add_bullet(doc, "关键设计思想: \"Errors are caught and returned as error text so the model can self-correct rather than crashing the loop\" (错误被捕获为文本返回，让模型能自我纠正而不是崩溃)")

    doc.add_page_break()

    # ── 3.2 factory.py ──
    doc.add_heading("3.2 factory.py — Agent 注册工厂 (159行)", level=2)

    doc.add_paragraph(
        "Agent 工厂模块负责创建和管理多个 Agent 实例。核心类是 AgentRegistry，"
        "提供注册 (register)、创建 (create_agent)、列表 (list_agents) 三个方法。"
        "文件底部定义了 create_default_registry() 工厂函数，返回预配置了 4 个 Agent 的注册表。"
    )

    doc.add_heading("3.2.1 AgentRegistry 类", level=3)
    add_table(doc,
        ["方法", "签名", "功能"],
        [
            ["register", "(name, system_prompt, tools, tool_map)", "注册一个 Agent 配置到内部字典"],
            ["create_agent", "(name, model) → BidAgent", "根据名称创建 BidAgent 实例"],
            ["list_agents", "() → List[str]", "返回所有已注册的 Agent 名称列表"],
        ]
    )

    doc.add_heading("3.2.2 四个预注册 Agent", level=3)
    add_table(doc,
        ["Agent Key", "中文名", "系统提示词键", "工具数量", "设计意图"],
        [
            ["supervisor", "智能助理", "SUPERVISOR_PROMPT (跨领域编排, 约30行中文)", "9个 (全部)", "跨域编排，整合标讯+企业+法规信息"],
            ["bidding_analyst", "标讯分析", "BIDDING_PROMPT (简短，1行中文)", "3个 (标讯)", "专注于招标公告搜索与趋势分析"],
            ["company_profiler", "企业画像", "COMPANY_PROMPT (简短，1行中文)", "3个 (企业)", "专注于企业信息查询与竞争分析"],
            ["legal_advisor", "法规咨询", "LEGAL_PROMPT (简短，1行中文)", "3个 (法规)", "专注于法律法规检索与法条查询"],
        ]
    )

    doc.add_heading("3.2.3 SUPERVISOR_PROMPT 详解", level=3)
    doc.add_paragraph(
        "Supervisor 的系统提示词是项目中最关键的一段提示工程 (Prompt Engineering)。"
        "它定义了以下内容："
    )
    add_bullet(doc, "工具分组：将9个工具按【标讯分析】【企业分析】【法规政策】三大类分组展示")
    add_bullet(doc, "工作流程：5步法 — 理解意图 → 按需调用工具 → 整合信息 → 引用具体数据 → 中文回复")
    add_bullet(doc, "示例场景：提供了3个跨域组合调用的具体例子，引导 LLM 进行工具编排")
    add_bullet(doc, "每个工具的 1 行中文描述，让 LLM 准确理解工具功能")

    doc.add_page_break()

    # ── 3.3 supervisor.py ──
    doc.add_heading("3.3 supervisor.py — 备用 Supervisor 构造器 (88行)", level=2)

    doc.add_paragraph(
        "这是一个备用模块，提供 create_supervisor_agent() 函数。"
        "它与 factory.py 中的 supervisor 注册项功能完全一致（都注册全部 9 个工具），"
        "但作为一个独立的函数暴露，方便在其他需要单独创建 Supervisor 的场景中使用。"
    )
    add_bullet(doc, "合并三个工具模块的 schema 列表: all_schemas = bidding_schema + company_schema + legal_schema")
    add_bullet(doc, "构建完整的 tool_handlers 字典，覆盖全部 9 个工具函数")
    add_bullet(doc, "复用 factory.py 中的 SUPERVISOR_PROMPT 作为系统提示词")
    add_bullet(doc, "存在意义：设计解耦 — 不依赖 AgentRegistry 也能获得 Supervisor 实例")

    doc.add_page_break()

    # ── 3.4 tools_bidding.py ──
    doc.add_heading("3.4 tools_bidding.py — 标讯分析工具集 (195行)", level=2)

    doc.add_paragraph(
        "该文件实现了 3 个标讯领域的工具函数，每个函数对应一个 OpenAI Function Calling schema 定义。"
        "所有工具函数操作 SQLite 数据库中的 bidding_notices 和 bidding_transactions 两张表。"
    )

    doc.add_heading("3.4.1 辅助函数", level=3)
    add_bullet(doc, "_get_conn(): 创建 SQLite 连接，设置 row_factory=sqlite3.Row 使查询结果可按列名访问")
    add_bullet(doc, "_parse_amount(raw): 解析中文金额格式，如 \"9,715.96万(元)\" → 97159.6 (注意这里只做数字提取和逗号移除，不进行万/亿单位换算，返回原始浮点数)")

    doc.add_heading("3.4.2 search_notices(keyword, limit=10)", level=3)
    doc.add_paragraph("按关键词在公告表中进行多列模糊搜索：")
    add_bullet(doc, "搜索列：notice_content (公告内容), notice_type (公告类型), tender_number (招标编号), successful_bidder (中标人), tendering_entity (招标单位)")
    add_bullet(doc, "SQL 模式：5 个 LIKE 条件 OR 连接 + LIMIT 限制")
    add_bullet(doc, "返回：字典列表，每行包含 notice_id, notice_type, tender_number, successful_bidder, awarded_amount, tendering_entity, tendering_address, notice_content")

    doc.add_heading("3.4.3 query_trends(industry, months=6)", level=3)
    doc.add_paragraph("统计特定行业近 N 个月的中标趋势：")
    add_bullet(doc, "数据来源：bidding_transactions 表")
    add_bullet(doc, "时间过滤：使用 datetime.strptime 解析 data 列为 YYYY/MM/DD 格式，计算 cutoff = now - months*30 天")
    add_bullet(doc, "聚合计算：total_amount (总中标金额), count (中标项目数), top_winners (TOP5 中标企业)")
    add_bullet(doc, "返回结构：{\"total_amount\": float, \"count\": int, \"top_winners\": [{\"name\": str, \"amount\": float}]}")

    doc.add_heading("3.4.4 get_notice_detail(notice_id)", level=3)
    doc.add_paragraph("根据公告 ID 查询完整内容：")
    add_bullet(doc, "SQL: SELECT * FROM bidding_notices WHERE notice_id = ?")
    add_bullet(doc, "未找到时返回: {\"error\": \"Notice {id} not found\"}")

    doc.add_heading("3.4.5 tools_schema 导出", level=3)
    doc.add_paragraph(
        "文件底部导出 tools_schema 列表，包含 3 个完整的 OpenAI Function Calling JSON Schema 定义。"
        "每个 schema 包含 type, function(name, description, parameters) 三部分，"
        "parameters 中定义了 properties（参数名、类型、描述、默认值）和 required（必选参数）。"
    )

    doc.add_page_break()

    # ── 3.5 tools_company.py ──
    doc.add_heading("3.5 tools_company.py — 企业画像工具集 (267行)", level=2)

    doc.add_paragraph(
        "该文件实现了 3 个企业分析工具函数，操作 qcc_companies (企查查企业数据) "
        "和 bidding_notices (中标记录) 两张表。"
    )

    doc.add_heading("3.5.1 _parse_capital(raw) — 中文注册资本解析", level=3)
    doc.add_paragraph(
        "这是企业工具集中最关键的辅助函数，负责将中文格式的注册资本转换为浮点数："
    )
    add_bullet(doc, "输入示例: \"9,715.96万(元)\" 或 \"1.5亿(元)\"")
    add_bullet(doc, "处理步骤: (1) 去除逗号 (2) 去除括号和中文字 (3) 正则提取首个数字 (4) 根据\"亿\"或\"万\"单位换算")
    add_bullet(doc, "换算规则: 含\"亿\"→ ×100,000,000; 含\"万\"或\"萬\"→ ×10,000")
    add_bullet(doc, "容错: 解析失败返回 0.0")

    doc.add_heading("3.5.2 search_companies(city, industry, min_capital, limit=20)", level=3)
    doc.add_paragraph("多条件企业筛选，支持三个维度组合查询：")
    add_bullet(doc, "city (城市): 对 city 列做 LIKE 模糊匹配")
    add_bullet(doc, "industry (行业): 对 industry 列做 LIKE 模糊匹配")
    add_bullet(doc, "min_capital (最低注册资本): 在 Python 层面过滤 (而非 SQL)，因为注册资本的数值需要 _parse_capital 解析")
    add_bullet(doc, "SQL 构建: 动态拼接 WHERE 条件，至少一个条件为真时用 AND 连接，无条件时用 WHERE 1=1")
    add_bullet(doc, "返回字段: company_name, legal_representative, registered_capital, establishment_date, province, city, district, industry, company_type, business_scope, registered_capital_value (解析后的数值)")

    doc.add_heading("3.5.3 get_company_profile(company_name)", level=3)
    doc.add_paragraph("企业全景画像，整合三类信息：")
    add_bullet(doc, "基本信息: 从 qcc_companies 模糊匹配企业名，返回全部字段 + 解析后的注册资本数值")
    add_bullet(doc, "中标记录: 在 bidding_notices 表中按 successful_bidder LIKE %company_name% 查询，LIMIT 20")
    add_bullet(doc, "同行统计: 在同行业中 COUNT(*) 统计企业总数 (peer_count)")
    add_bullet(doc, "返回结构: {基本信息字段..., \"bidding_records\": [...], \"peer_count\": int}")

    doc.add_heading("3.5.4 find_competitors(company_name, limit=10)", level=3)
    doc.add_paragraph("竞争对手分析，按注册资本降序排列：")
    add_bullet(doc, "Step 1: 查询目标企业的 city 和 industry")
    add_bullet(doc, "Step 2: 在同城 + 同行业的条件下查询所有企业 (排除自身)")
    add_bullet(doc, "Step 3: 解析所有企业的注册资本数值")
    add_bullet(doc, "Step 4: 按 registered_capital_value 降序排列，取 TOP N")

    doc.add_page_break()

    # ── 3.6 tools_legal.py ──
    doc.add_heading("3.6 tools_legal.py — 法规检索工具集 (191行)", level=2)

    doc.add_paragraph(
        "该文件实现了 3 个法律法规检索工具，结合了 SQLite 全文搜索和 TF-IDF 语义搜索两种策略。"
        "文件级别的 _vec_store 变量实现了向量存储的单例惰性加载。"
    )

    doc.add_heading("3.6.1 semantic_search_laws(query, top_k=5)", level=3)
    doc.add_paragraph("TF-IDF 语义搜索，是法规检索的推荐优先方式：")
    add_bullet(doc, "调用 _get_vec_store().search(query, top_k) 执行 TF-IDF + Cosine Similarity 搜索")
    add_bullet(doc, "索引未构建时返回错误提示，引导用户先运行 build_index.py")
    add_bullet(doc, "适合自然语言查询，如 \"投标保证金有什么规定\"、\"中小企业优惠政策\"")

    doc.add_heading("3.6.2 search_laws(keyword, limit=5)", level=3)
    doc.add_paragraph("SQLite 全文关键词搜索：")
    add_bullet(doc, "在 laws 表的 content 和 document_title 两列中进行 LIKE 模糊匹配")
    add_bullet(doc, "返回字段: document_type, document_title, chapter_title, section_title, article_number, content")

    doc.add_heading("3.6.3 get_article(law_title, article_number)", level=3)
    doc.add_paragraph("精确法条原文查询，使用三层回退策略：")
    add_bullet(doc, "第1层 — 精确匹配: document_title LIKE + article_number = (精确匹配条款编号)")
    add_bullet(doc, "第2层 — 条款模糊匹配: document_title LIKE + article_number LIKE (条款编号作为子串)")
    add_bullet(doc, "第3层 — 回退: 仅按法律名称匹配，返回该法律的任意一条 (last resort)")
    add_bullet(doc, "全部失败: 返回 {\"error\": \"Article not found\"}")

    doc.add_page_break()

    # ── 3.7 vector_store.py ──
    doc.add_heading("3.7 vector_store.py — TF-IDF 向量存储器 (174行)", level=2)

    doc.add_paragraph(
        "基于 scikit-learn TfidfVectorizer 实现的中文文档语义检索引擎。"
        "设计原则是离线构建、在线加载、可替换（接口设计允许将来替换为神经网络 Embedding）。"
    )

    doc.add_heading("3.7.1 索引文件", level=3)
    add_table(doc,
        ["文件", "内容", "格式"],
        [
            ["tfidf_vectorizer.pkl", "TfidfVectorizer 实例", "pickle 序列化"],
            ["tfidf_matrix.npy", "TF-IDF 矩阵 (docs × features)", "NumPy 二进制"],
            ["laws_policy_meta.json", "文档元数据列表", "JSON UTF-8"],
        ]
    )

    doc.add_heading("3.7.2 关键参数", level=3)
    add_bullet(doc, "analyzer='char_wb': 字符级 n-gram (word boundary)，专门针对中文无空格分词的特点设计。char_wb 从词边界内部生成字符 n-gram，能捕获中文词汇的局部模式")
    add_bullet(doc, "max_features=5000: 限制特征维度，控制矩阵大小")
    add_bullet(doc, "ngram_range=(1, 2): 同时使用 unigram 和 bigram，捕捉单字和双字特征")

    doc.add_heading("3.7.3 核心方法", level=3)
    add_bullet(doc, "search(query, top_k=5): 将查询文本转换为 TF-IDF 向量 → 计算与矩阵的 cosine_similarity → 按分数降序排列 → 过滤 score<=0 的结果 → 返回 top_k 结果 (附加 _score 字段)")
    add_bullet(doc, "build_from_db(db_path, output_dir) [静态方法]: 从 SQLite 读取 laws + t_policy + policy 三张表 → 拼接文本字段 → TF-IDF 拟合 → 持久化到磁盘")
    add_bullet(doc, "文本拼接策略: document_title + chapter_title + article_number + content[:2000] (截断长文本控制特征空间)")
    add_bullet(doc, "策略表兼容: 适用于 t_policy 和 policy 两种表名，支持 title/policy_title 和 text/content 两种列名变体")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 4: 展示与 API 层
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("四、展示与 API 层", level=1)

    # ── 4.1 api.py ──
    doc.add_heading("4.1 api.py — FastAPI 后端服务 (127行)", level=2)

    doc.add_paragraph(
        "基于 FastAPI 框架的 RESTful API 服务，为前端或其他系统提供标准化的 Agent 调用接口。"
        "启动命令: uvicorn api:app --reload --port 8000"
    )

    doc.add_heading("4.1.1 应用配置", level=3)
    add_bullet(doc, "FastAPI 实例: title=\"招投标智能体平台 API\", version=\"1.0.0\"")
    add_bullet(doc, "CORS 中间件: 允许所有来源 (allow_origins=[\"*\"])，支持跨域访问")
    add_bullet(doc, "Agent 注册表: 全局单例 registry = create_default_registry()")
    add_bullet(doc, "会话存储: sessions: Dict[str, List[Dict]] — 内存字典，key=session_id, value=messages 列表")

    doc.add_heading("4.1.2 Pydantic 数据模型", level=3)
    add_table(doc,
        ["模型", "字段", "用途"],
        [
            ["ChatRequest", "agent_name (默认supervisor), message (必填), session_id (可选), model", "POST /chat 请求体"],
            ["ChatResponse", "session_id, agent_name, response, model", "POST /chat 响应体"],
            ["AgentInfo", "name, tools_count", "GET /agents 响应元素"],
            ["AgentListResponse", "agents: List[AgentInfo]", "GET /agents 响应体"],
        ]
    )

    doc.add_heading("4.1.3 API 端点", level=3)
    add_table(doc,
        ["方法", "路径", "功能", "关键细节"],
        [
            ["GET", "/", "服务信息", "返回 service 名称、版本、文档链接 /docs"],
            ["GET", "/agents", "Agent 列表", "从 registry 读取，包含 name 和 tools_count"],
            ["POST", "/chat", "发送消息", "核心端点。生成/使用 session_id 维护对话连续性。异常捕获为 HTTP 500"],
            ["GET", "/sessions/{session_id}", "会话历史", "返回指定会话的全部 messages。不存在返回 404"],
            ["GET", "/health", "健康检查", "返回 {\"status\": \"ok\"}"],
        ]
    )

    doc.add_heading("4.1.4 会话管理机制", level=3)
    doc.add_paragraph(
        "使用内存字典 sessions 存储对话历史。session_id 生成策略: "
        "由客户端传入 (用于持续对话) 或服务端自动生成 uuid4 前 8 位。"
        "每次 /chat 调用后，将 user message 和 assistant response 追加到 sessions[session_id]。"
        "注意：服务重启后所有会话丢失（内存存储）。"
    )

    doc.add_page_break()

    # ── 4.2 app.py ──
    doc.add_heading("4.2 app.py — Streamlit 前端界面 (106行)", level=2)

    doc.add_paragraph(
        "基于 Streamlit 框架的多 Agent 对话 Web 界面。"
        "启动命令: streamlit run app.py"
    )

    doc.add_heading("4.2.1 UI 布局", level=3)
    add_bullet(doc, "页面配置: page_title=\"招投标智能体平台\", layout=\"wide\" (宽屏布局)")
    add_bullet(doc, "侧边栏: Agent 选择下拉框 (4个选项+Emoji图标)、清空历史复选框、模型名称显示")
    add_bullet(doc, "主区域: 平台标题 → 当前 Agent 名称 → 对话消息列表 → 聊天输入框")

    doc.add_heading("4.2.2 Agent 选项映射", level=3)
    add_bullet(doc, "\"🤖 智能助理\" → \"supervisor\"")
    add_bullet(doc, "\"📊 标讯分析\" → \"bidding_analyst\"")
    add_bullet(doc, "\"🏢 企业画像\" → \"company_profiler\"")
    add_bullet(doc, "\"⚖️ 法规咨询\" → \"legal_advisor\"")

    doc.add_heading("4.2.3 Session State 管理", level=3)
    doc.add_paragraph(
        "使用 Streamlit 的 st.session_state 管理前端状态，核心数据结构："
    )
    add_bullet(doc, "histories: Dict[str, List[Dict]] — 每个 Agent 维护独立的对话历史，切换 Agent 不会丢失之前的对话")
    add_bullet(doc, "current_agent: str — 当前选中的 Agent key")
    add_bullet(doc, "切换 Agent 行为: 当 \"切换Agent时清空对话历史\" 复选框选中时，会将目标 Agent 的历史置空")

    doc.add_heading("4.2.4 对话流程", level=3)
    add_bullet(doc, "用户在聊天框输入问题 → 追加到当前 Agent 的对话历史")
    add_bullet(doc, "显示用户消息 (st.chat_message(\"user\"))")
    add_bullet(doc, "显示加载动画 (st.spinner(\"分析中...\"))")
    add_bullet(doc, "调用 agent.run(user_input) 获取回复")
    add_bullet(doc, "显示助手回复 (st.chat_message(\"assistant\"))")
    add_bullet(doc, "异常处理: 显示 \"❌ 调用失败：{error}\"")

    doc.add_heading("4.2.5 缓存策略", level=3)
    add_bullet(doc, "@st.cache_resource: 缓存 AgentRegistry 实例，避免每次对话都重新创建")
    add_bullet(doc, "注意: BidAgent 实例本身不缓存，每次对话都新建（因为每个请求独立执行 Tool Calling 循环）")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 5: 脚本工具层
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("五、脚本工具层 (scripts/)", level=1)

    # ── 5.1 init_db.py ──
    doc.add_heading("5.1 init_db.py — 数据库初始化管道 (315行)", level=2)

    doc.add_paragraph(
        "这是一个全自动化的 CSV → SQLite 数据导入管道，是整个项目数据层的基石。"
        "运行: python scripts/init_db.py。输出: data/bid_agent.db。"
    )

    doc.add_heading("5.1.1 核心设计", level=3)
    add_bullet(doc, "自动编码检测: 依次尝试 utf-8 → gbk → gb18030 → gb2312 → latin-1 编码读取 CSV")
    add_bullet(doc, "中文列名标准化: 维护 CN_COLUMN_MAP 字典 (约60个映射项)，将中文列名（如\"企业名称\"）转换为英文（company_name）")
    add_bullet(doc, "特殊文件名处理: FILENAME_MAP 将 \"招投标交易\" 映射为 \"bidding_transactions\"")
    add_bullet(doc, "SQLite 类型推断: 根据 pandas Series dtype 推断 SQLite 类型 (INTEGER/REAL/TEXT)")

    doc.add_heading("5.1.2 处理流程", level=3)
    add_bullet(doc, "Step 1: 扫描 data/ 目录下所有 .csv 文件")
    add_bullet(doc, "Step 2: 删除旧的 bid_agent.db (全新构建，保证幂等)")
    add_bullet(doc, "Step 3: 对每个 CSV 文件执行 import_csv()：")
    add_bullet(doc, "  - sanitize_table_name(): 文件名转 SQLite 安全表名", level=1)
    add_bullet(doc, "  - detect_encoding(): 自动检测文件编码", level=1)
    add_bullet(doc, "  - pd.read_csv(): 使用检测到的编码读取，keep_default_na=False (不将空字符串转为 NaN)", level=1)
    add_bullet(doc, "  - standardize_columns(): 中文列名 → 英文列名 + 安全化处理", level=1)
    add_bullet(doc, "  - get_or_create_table(): 建表 (DDL)，如果已存在则跳过", level=1)
    add_bullet(doc, "  - INSERT OR REPLACE: 全量写入数据 (幂等设计)", level=1)
    add_bullet(doc, "  - create_indexes(): 自动为 id/name/title/date/code/province/city 等常见查询列创建索引", level=1)
    add_bullet(doc, "Step 4: 输出导入汇总 (每个文件的行数 + 总行数 + 数据库表列表)")

    doc.add_heading("5.1.3 表名安全化 (sanitize_table_name)", level=3)
    add_bullet(doc, "去除非字母数字下划线字符")
    add_bullet(doc, "合并连续下划线")
    add_bullet(doc, "全部转小写")
    add_bullet(doc, "数字开头自动加 \"t_\" 前缀 (符合 SQL 标识符规范)")

    doc.add_page_break()

    # ── 5.2 build_index.py ──
    doc.add_heading("5.2 build_index.py — 向量索引构建 (29行)", level=2)

    doc.add_paragraph(
        "简洁的索引构建入口脚本，封装 VectorStore.build_from_db() 的调用。"
    )
    add_bullet(doc, "从 SQLite 数据库 (DB_PATH) 中读取法律和政策数据")
    add_bullet(doc, "调用 VectorStore.build_from_db() 构建 TF-IDF 索引")
    add_bullet(doc, "索引输出到 data/vector_index/ 目录")
    add_bullet(doc, "自动验证: 用 \"招标投标保证金\" 查询 top 3 结果并打印分数")

    doc.add_page_break()

    # ── 5.3 evaluate.py ──
    doc.add_heading("5.3 evaluate.py — 自动化评测框架 (229行)", level=2)

    doc.add_paragraph(
        "完整的 Agent 评测框架，用于量化评估 Agent 的工具调用准确率和响应质量。"
    )

    doc.add_heading("5.3.1 EvalResult 类", level=3)
    doc.add_paragraph("封装单条评测用例的执行结果，包含以下字段：")
    add_bullet(doc, "case_id, question, domain: 用例标识")
    add_bullet(doc, "tool_called: 是否调用了至少一个工具")
    add_bullet(doc, "tools_used: 实际调用的工具名列表")
    add_bullet(doc, "correct_tool_called: 是否调用了预期的工具")
    add_bullet(doc, "has_keywords: 响应中是否包含预期关键词")
    add_bullet(doc, "response, elapsed, error: 响应内容、耗时、异常信息")
    add_bullet(doc, "passed 属性: tool_called AND correct_tool_called AND has_keywords")

    doc.add_heading("5.3.2 工具调用追踪机制", level=3)
    doc.add_paragraph(
        "使用 monkey-patch 技术在运行时拦截工具调用："
    )
    add_bullet(doc, "在 run_eval() 中: 保存 agent._execute_tool 引用 → 替换为 tracking_execute → 执行 run() → 恢复原函数")
    add_bullet(doc, "tracking_execute 在执行前将工具名追加到 tools_called 列表，再调用原始函数")
    add_bullet(doc, "这种无侵入式的 instrumentation 设计避免了修改 BidAgent 源代码")

    doc.add_heading("5.3.3 评测指标", level=3)
    add_table(doc,
        ["指标", "计算方式", "含义"],
        [
            ["工具调用率", "调用至少1个工具的用例数 / 总用例数", "Agent 是否主动使用工具"],
            ["工具选择准确率", "调用预期工具的用例数 / 总用例数", "Agent 是否正确选择了工具"],
            ["关键词覆盖率", "包含预期关键词的用例数 / 总用例数", "响应内容是否相关"],
            ["平均响应时间", "所有用例耗时均值", "系统性能指标"],
            ["通过率", "三项指标全部满足的用例比例", "综合质量指标"],
        ]
    )

    doc.add_heading("5.3.4 评测流程", level=3)
    add_bullet(doc, "加载 tests/eval_data.json (50条标注用例)")
    add_bullet(doc, "对 4 个 Agent (supervisor, bidding_analyst, company_profiler, legal_advisor) 分别运行全部 50 条用例")
    add_bullet(doc, "每个 Agent 生成独立的详细结果文件: tests/eval_result_{agent_name}.json")
    add_bullet(doc, "按领域维度输出分类统计")
    add_bullet(doc, "列出失败用例及原因")

    doc.add_page_break()

    # ── 5.4 generate_report.py ──
    doc.add_heading("5.4 generate_report.py — Word 技术报告生成器 (184行)", level=2)

    doc.add_paragraph(
        "从 scripts/report_content.json 读取结构化内容，使用 python-docx 生成格式化的 Word 技术报告。"
    )
    add_bullet(doc, "封面生成: 标题 + 副标题 + 标签语 + 元信息，居中对齐")
    add_bullet(doc, "10 个章节: 项目概述 → 架构 → 搭建流程 → 数据流 → 数据库设计 → 评测体系 → 测试覆盖 → 文件清单 → 启动方式 → 简历亮点")
    add_bullet(doc, "自动统计: 遍历文件清单计算每个文件的实际行数，汇总总行数")
    add_bullet(doc, "样式: 标题蓝色 (#1A56DB)，代码块 Consolas 9pt，表格 Light Shading Accent 1 样式")
    add_bullet(doc, "输出: 招投标智能体平台_技术报告.docx")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 6: 测试与评测体系
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("六、测试与评测体系 (tests/)", level=1)

    doc.add_heading("6.1 conftest.py — pytest 配置 (5行)", level=2)
    doc.add_paragraph(
        "将项目根目录添加到 sys.path，确保测试文件中可以 import agent 包。"
    )

    doc.add_heading("6.2 test_agent.py — BidAgent 核心单元测试 (334行)", level=2)
    doc.add_paragraph("使用 unittest.mock.patch 完全模拟 API 调用，不依赖真实网络或数据库。")

    add_table(doc,
        ["测试类", "用例", "验证点"],
        [
            ["TestAgentNoTools", "test_agent_returns_direct_text", "无工具时直接返回 LLM 文本"],
            ["TestAgentNoTools", "test_agent_includes_system_prompt", "system prompt 被正确传递到 API 调用"],
            ["TestAgentWithTool", "test_agent_executes_tool_and_returns_result", "工具调用 → 执行 → 返回结果 完整链路"],
            ["TestAgentWithTool", "test_agent_handles_tool_execution_error_as_content", "工具执行异常被捕获为文本，循环继续"],
            ["TestAgentWithTool", "test_agent_unknown_tool_returns_error_message", "未知工具不崩溃，返回错误信息"],
            ["TestAgentWithTool", "test_agent_max_turns_exceeded_returns_timeout", "超过 10 轮返回\"处理超时\""],
            ["TestAgentApiRetry", "test_agent_retries_on_api_failure", "前2次失败第3次成功 → 重试生效"],
            ["TestAgentApiRetry", "test_agent_raises_after_max_retries", "3次全失败 → 抛出 RuntimeError"],
        ]
    )

    doc.add_heading("6.3 test_tools.py — 工具函数集成测试 (124行)", level=2)
    doc.add_paragraph("对真实 SQLite 数据库运行，验证工具函数的数据访问正确性。")

    add_table(doc,
        ["测试类", "用例数", "覆盖工具"],
        [
            ["TestBiddingTools", "5个", "search_notices, query_trends, get_notice_detail"],
            ["TestCompanyTools", "7个", "search_companies, get_company_profile, find_competitors"],
            ["TestLegalTools", "4个", "search_laws, get_article"],
        ]
    )
    doc.add_paragraph(
        "关键测试策略: (1) 验证返回类型和字段完整性 (2) 验证查询条件的过滤效果 "
        "(3) 验证边界情况（无结果、不存在的ID/名称） (4) 验证排序逻辑（竞争对手按注册资本降序）"
        "(5) 先查询获取真实数据再将其作为后续测试的输入（如先 search_companies 再 get_company_profile）"
    )

    doc.add_heading("6.4 eval_data.json — 标注评测数据集", level=2)
    doc.add_paragraph(
        "50 条人工标注的 QA 评测用例，是衡量系统质量的标准基准："
    )
    add_bullet(doc, "标讯领域 (B01-B12): 12 条 — 招标公告搜索、趋势统计、详情查询")
    add_bullet(doc, "企业查询 (C01-C15): 15 条 — 企业筛选、画像、竞争对手")
    add_bullet(doc, "法规检索 (L01-L10): 10 条 — 法律法规搜索、法条原文查询")
    add_bullet(doc, "跨域综合 (X01-X13): 13 条 — 需要跨领域工具编排的复杂查询")
    add_bullet(doc, "每条包含: id, question, domain, expected_tools, expected_keywords, min_results")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 7: 数据层
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("七、数据层 (data/)", level=1)

    doc.add_heading("7.1 数据文件清单", level=2)
    add_table(doc,
        ["CSV 文件", "对应数据库表", "主要内容"],
        [
            ["bidding_notices.csv", "bidding_notices", "招标公告: 编号、类型、中标人、金额、内容"],
            ["招投标交易.csv", "bidding_transactions", "中标交易: 中标方、金额、日期、类型、内容"],
            ["companies.csv", "companies", "企业基本信息（备用数据源）"],
            ["company_info.csv", "company_info", "企业补充信息"],
            ["company_info1.csv", "company_info1", "企业补充信息（第二批次）"],
            ["qcc_companies.csv", "qcc_companies", "企查查企业数据: 名称、法人、注册资本、行业、城市等"],
            ["ods_company_detail.csv", "ods_company_detail", "ODS 层企业详情"],
            ["laws.csv", "laws", "法律法规: 文档类型/标题/章节/条款编号/内容"],
            ["policy.csv", "policy", "政策文件"],
            ["policy_metadata.csv", "policy_metadata", "政策元数据"],
            ["t_policy.csv", "t_policy", "政策文件（第二数据源）"],
        ]
    )

    doc.add_heading("7.2 核心数据表", level=2)

    doc.add_heading("bidding_notices (招标公告表)", level=3)
    doc.add_paragraph(
        "主键: notice_id。核心字段: notice_type (公告类型), tender_number (招标编号), "
        "successful_bidder (中标人), awarded_amount (中标金额), tendering_entity (招标单位), "
        "tendering_address (招标地址), notice_content (公告内容全文)。"
        "该表是标讯分析工具 (search_notices, get_notice_detail) 和企业画像工具 "
        "(get_company_profile 的中标记录部分) 的数据来源。"
    )

    doc.add_heading("bidding_transactions (中标交易表)", level=3)
    doc.add_paragraph(
        "核心字段: bidding (中标方/投标方), bid_amount (投标金额), data (日期), "
        "content (交易内容), type (交易类型)。"
        "该表是 query_trends 工具的数据来源。日期格式为 YYYY/MM/DD。"
    )

    doc.add_heading("qcc_companies (企业信息表)", level=3)
    doc.add_paragraph(
        "核心字段: company_name, legal_representative (法定代表人), registered_capital (注册资本), "
        "establishment_date (成立日期), province, city, district (省/市/区县), "
        "industry (所属行业), company_type (企业类型), business_scope (经营范围), "
        "unified_social_credit_code (统一社会信用代码), insured_count (参保人数)。"
        "该表是企业画像工具的主要数据来源。"
    )

    doc.add_heading("laws (法律法规表)", level=3)
    doc.add_paragraph(
        "核心字段: document_type (文档类型，如\"法律\"/\"行政法规\"), document_title (文档标题，如\"中华人民共和国招标投标法\"), "
        "chapter_title (章节标题), section_title (节标题), article_number (条款编号), content (条款原文)。"
        "该表是法规检索工具和 TF-IDF 索引的数据来源。"
    )

    doc.add_heading("7.3 向量索引 (vector_index/)", level=2)
    doc.add_paragraph(
        "3 个文件组成的 TF-IDF 检索引擎持久化数据: tfidf_vectorizer.pkl (TfidfVectorizer 实例，"
        "约 5000 特征 × char_wb 分析器), tfidf_matrix.npy (文档-词项矩阵), "
        "laws_policy_meta.json (每个文档的 source, title, text 元数据)。"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 8: 配置文件
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("八、配置文件与环境变量", level=1)

    doc.add_heading("8.1 .env 环境变量", level=2)
    add_table(doc,
        ["变量名", "值", "用途"],
        [
            ["DEEPSEEK_API_KEY", "your-deepseek-api-key-here", "DeepSeek API 认证密钥"],
            ["DEEPSEEK_BASE_URL", "https://api.deepseek.com", "API 基础 URL (兼容 OpenAI SDK)"],
            ["DB_PATH", "data/bid_agent.db", "SQLite 数据库文件路径"],
            ["LOG_LEVEL", "INFO", "Python logging 日志级别"],
        ]
    )
    doc.add_paragraph(
        "加载方式: api.py 和 app.py 都在文件开头调用 load_dotenv() 加载 .env 文件。"
        "BidAgent 构造函数中从 os.getenv() 读取 API 配置。"
    )

    doc.add_heading("8.2 requirements.txt 依赖清单", level=2)
    add_table(doc,
        ["包名", "最低版本", "用途"],
        [
            ["openai", ">=1.0.0", "OpenAI 兼容 API 客户端 (用于调用 DeepSeek)"],
            ["streamlit", ">=1.32.0", "Web 前端框架"],
            ["python-dotenv", ">=1.0.0", "环境变量管理"],
            ["pandas", ">=2.0.0", "数据处理与 CSV 导入"],
            ["pytest", ">=8.0.0", "测试框架"],
            ["scikit-learn", ">=1.0.0", "TF-IDF 向量化和余弦相似度"],
            ["fastapi", ">=0.100.0", "REST API 框架"],
            ["uvicorn", ">=0.23.0", "ASGI 服务器 (运行 FastAPI)"],
        ]
    )

    doc.add_heading("8.3 CLAUDE.md 项目规范", level=2)
    doc.add_paragraph(
        "定义了 AI 辅助开发的编码规范: Python 3.11+、pip 包管理、函数必须添加 type hints、"
        "公共函数使用 docstring、单文件不超过 300 行、使用 SQLite 数据库、"
        "OpenAI SDK function calling 模式、兼容 DeepSeek API、默认模型 deepseek-chat、"
        ".env 管理环境变量、Streamlit 构建前端、logging 记录日志、pytest 测试。"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 9: 完整请求链路
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("九、完整请求链路剖析", level=1)

    doc.add_paragraph(
        "以下追踪一个完整的用户对话请求从输入到返回的全链路。"
    )

    doc.add_heading("9.1 Streamlit 前端路径", level=2)
    add_code_block(doc, "用户输入: \"合肥建筑行业有哪些中标企业？\"")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "app.py: st.chat_input() 捕获输入")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "app.py: registry.create_agent(agent_key) → BidAgent 实例化")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "app.py: agent.run(user_input)")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: 构建 messages = [system_prompt, user_input]")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: Round 1 — _call_api(messages, tools) → DeepSeek API")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "LLM 返回: finish_reason=\"tool_calls\", 要求调用 search_companies(\"合肥\", \"建筑\")")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: _execute_tool(\"search_companies\", '{\"city\": \"合肥\", \"industry\": \"建筑\"}')")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "tools_company.py: SQLite 查询 → 返回企业列表")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: 工具结果追加到 messages，继续 Round 2")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "LLM 返回: finish_reason=\"tool_calls\", 要求调用 search_notices(\"建筑\")")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: _execute_tool(\"search_notices\", ...) → 工具执行 → 追加结果")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: Round 3 — LLM 整合两轮结果")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "LLM 返回: finish_reason=\"stop\", content=\"根据分析，合肥建筑行业...\"")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "core.py: 返回文本 response")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "app.py: st.markdown(response) 渲染到界面")

    doc.add_heading("9.2 FastAPI 后端路径", level=2)
    doc.add_paragraph(
        "当通过 API 访问时，流程等效但增加了 HTTP 层："
    )
    add_code_block(doc, "POST /chat {\"agent_name\": \"supervisor\", \"message\": \"...\"}")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "api.py: 验证 agent_name 有效性 → 生成 session_id")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "api.py: registry.create_agent(name) → agent.run(message)")
    add_code_block(doc, "    ↓")
    add_code_block(doc, "api.py: 存入 sessions[session_id] → 返回 ChatResponse")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 10: 设计模式
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("十、设计模式与架构决策", level=1)

    doc.add_heading("10.1 Supervisor + Specialists 编排模式", level=2)
    doc.add_paragraph(
        "核心架构决策：不是让一个 Agent 承担所有任务，也不是让多个 Agent 互相通信，"
        "而是让一个 Supervisor 拥有全部工具的访问权，由 LLM 的 Function Calling 能力自动编排。"
        "Specialist Agent 存在的主要价值是：(1) 在领域内聚焦时成本更低 (2) 作为对比基准验证 Supervisor 的价值。"
    )

    doc.add_heading("10.2 Tool Calling 循环", level=2)
    doc.add_paragraph(
        "采用标准 ReAct (Reasoning + Acting) 模式: LLM 推理 → 决定调用工具 → 工具执行 → "
        "结果返回 → 继续推理。循环由 BidAgent.run() 中的 for 循环 + finish_reason 判断驱动。"
        "关键安全机制: MAX_ROUNDS=10 防止无限循环（如 LLM 不断调用工具但不给出最终答案）。"
    )

    doc.add_heading("10.3 优雅降级 (Graceful Degradation)", level=2)
    add_bullet(doc, "API 失败: 指数退避重试 3 次 → 彻底失败抛出异常 → 上层 (api.py/app.py) 捕获显示错误")
    add_bullet(doc, "工具执行异常: 被捕获为文本返回给 LLM (不崩溃循环) → LLM 根据错误信息尝试其他策略")
    add_bullet(doc, "未知工具: 返回 \"Unknown tool\" 文本 → LLM 尝试其他工具")
    add_bullet(doc, "法条未找到: get_article 三层回退 → 高层级失败不会导致整个请求失败")
    add_bullet(doc, "向量索引未构建: semantic_search_laws 返回友好错误提示")

    doc.add_heading("10.4 工厂模式 (AgentRegistry)", level=2)
    doc.add_paragraph(
        "AgentRegistry 封装了 Agent 的创建逻辑：(1) 集中管理所有 Agent 的配置 (register/register) "
        "(2) 提供统一的创建接口 (create_agent) (3) 支持动态查询 (list_agents)。"
        "create_default_registry() 工厂函数预配置了 4 个 Agent，保证系统开箱即用。"
    )

    doc.add_heading("10.5 单例惰性加载 (VectorStore)", level=2)
    doc.add_paragraph(
        "VectorStore 采用惰性加载策略：构造函数只记录索引路径，真正的文件加载延迟到首次 search() 调用。"
        "tools_legal.py 中的 _get_vec_store() 函数使用模块级全局变量实现单例模式。"
    )

    doc.add_heading("10.6 接口预留设计", level=2)
    doc.add_paragraph(
        "VectorStore 类的文档明确说明：\"Interface is designed so neural embeddings can be swapped in later.\""
        "search(query, top_k) → List[Dict] 这个接口签名与任何嵌入模型的搜索接口兼容，"
        "未来可以将 TF-IDF 替换为 Sentence-BERT 或 OpenAI Embeddings 而不影响调用方。"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 11: 关键技术细节
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("十一、关键技术细节与源码级分析", level=1)

    doc.add_heading("11.1 为什么使用 char_wb 分析器", level=2)
    doc.add_paragraph(
        "中文文本没有空格分词，传统的 word 级别分析器对中文效果很差。scikit-learn 提供三种分析器："
    )
    add_bullet(doc, "word: 按空格/标点分词 → 中文几乎不可用")
    add_bullet(doc, "char: 纯字符级 n-gram → 会产生大量无意义的特征（如标点符号的 n-gram）")
    add_bullet(doc, "char_wb (word boundary): 仅在词边界内部生成字符 n-gram → 对中文而言，每个汉字被视为一个'词'的边界，n-gram 在连续的汉字序列中生成 → 既能捕获 \"投标\"、\"招标\" 这样的双字词，也能捕获 \"保证\"、\"保证金\" 这样的多字组合")
    doc.add_paragraph(
        "配合 ngram_range=(1,2)，可以同时匹配单字 (unigram) 和双字组合 (bigram)，"
        "对中文法律文本的语义搜索效果显著优于纯关键词匹配。"
    )

    doc.add_heading("11.2 中文金额解析的复杂性", level=2)
    doc.add_paragraph(
        "项目中有两个不同的金额解析函数，它们的行为有微妙差异："
    )
    add_bullet(doc, "tools_bidding._parse_amount(): 只去除中文字符和逗号，不做单位换算。用于中标金额统计（本身已经是统一单位）")
    add_bullet(doc, "tools_company._parse_capital(): 完整的 \"亿\"/\"万\" 单位换算 (×100,000,000 / ×10,000)。用于注册资本筛选（需要跨单位比较）")
    doc.add_paragraph(
        "这种差异化设计反映了实际数据的特点：中标金额数据可能已经是统一数值格式，"
        "而注册资本数据混合使用 \"万\" 和 \"亿\" 单位，需要标准化后才能进行大小比较。"
    )

    doc.add_heading("11.3 SQLite row_factory 的妙用", level=2)
    doc.add_paragraph(
        "每个工具的 _get_conn() 函数都设置 conn.row_factory = sqlite3.Row，"
        "这使得查询结果可以同时通过索引 (row[0]) 和列名 (row[\"company_name\"]) 访问。"
        "配合 dict(r) 转换，可以方便地返回 JSON 兼容的字典列表，直接用于 API 响应和 LLM 上下文。"
    )

    doc.add_heading("11.4 DeepSeek API 的 OpenAI 兼容性", level=2)
    doc.add_paragraph(
        "项目使用 openai 官方 SDK 但连接的是 DeepSeek API。实现兼容的关键："
    )
    add_bullet(doc, "base_url 指向 https://api.deepseek.com (而非 https://api.openai.com)")
    add_bullet(doc, "model=\"deepseek-chat\" (DeepSeek 的模型标识)")
    add_bullet(doc, "Function Calling 的 JSON Schema 格式完全遵循 OpenAI 规范 (type/function/name/description/parameters)")
    add_bullet(doc, "Tool Calling 的执行流程 (finish_reason=\"tool_calls\" / \"stop\") 与 OpenAI 一致")
    add_bullet(doc, "这种兼容设计使得将来切换到其他兼容 OpenAI 的 API (如 Moonshot, 智谱 GLM, 通义千问) 几乎零成本")

    doc.add_heading("11.5 幂等性设计", level=2)
    doc.add_paragraph(
        "init_db.py 使用 INSERT OR REPLACE 语句，配合先删除旧数据库的策略，保证每次运行产生确定性的干净数据库。"
        "build_index.py 每次运行覆盖旧的向量索引文件。这种设计避免了增量更新带来的数据一致性问题和复杂度。"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 12: 评测结果
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("十二、评测结果与性能分析", level=1)

    doc.add_heading("12.1 各 Agent 评测对比", level=2)
    add_table(doc,
        ["Agent", "通过率", "工具准确率", "关键词覆盖率", "平均响应时间"],
        [
            ["supervisor (智能助理)", "86.0%", "88.0%", "96.0%", "13.0s"],
            ["bidding_analyst (标讯分析)", "42.0%", "46.0%", "94.0%", "11.2s"],
            ["company_profiler (企业画像)", "34.0%", "34.0%", "98.0%", "9.9s"],
            ["legal_advisor (法规咨询)", "32.0%", "34.0%", "92.0%", "12.1s"],
        ]
    )

    doc.add_heading("12.2 关键洞察", level=2)
    add_bullet(doc, "Supervisor 优势显著: 通过率 (86%) 是各 Specialist (32%-42%) 的 2-2.7 倍。核心原因是 Supervisor 拥有全部工具，在 13 条跨域用例 (X01-X13) 中表现优异，而 Specialist 在跨域场景下工具不足")
    add_bullet(doc, "工具选择准确率差异: Supervisor 的 88% (9选N) vs Specialist 的 34%-46% (3选N)。即使 Supervisor 拥有更多工具需要选择，其准确率反而更高，说明丰富的系统提示词和工具描述对准确性有决定性影响")
    add_bullet(doc, "关键词覆盖率整体较高 (92%-98%): 说明当 LLM 生成回复时，内容质量是及格的。瓶颈在工具调用正确性而非文本生成质量")
    add_bullet(doc, "响应时间合理: 9.9s-13.0s 的范围对于需要多轮工具调用的 AI Agent 来说是正常的（包含 LLM API 网络延迟 + 数据库查询 + 多轮循环）")

    doc.add_heading("12.3 Specialist 低通过率分析", level=2)
    doc.add_paragraph(
        "bidding_analyst、company_profiler、legal_advisor 的通过率偏低 (32%-42%)，"
        "主要原因是评测数据集包含 13 条跨域综合用例 (X01-X13)，而这些 Specialist 只拥有 3 个单领域工具，"
        "无法完成跨域任务。这是设计上的合理取舍——Specialist 的设计目标就是单领域高效，而不是全场景覆盖。"
        "在实际使用中，用户可以通过选择 Supervisor 来处理复杂查询，或选择 Specialist 进行单一领域的深度分析。"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 13: 启动指南
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("十三、启动与运行指南", level=1)

    doc.add_heading("13.1 环境准备", level=2)
    add_code_block(doc, "# 1. 安装 Python 依赖")
    add_code_block(doc, "pip install -r requirements.txt")
    add_code_block(doc, "")
    add_code_block(doc, "# 2. 配置环境变量")
    add_code_block(doc, "cp .env.example .env")
    add_code_block(doc, "# 编辑 .env 文件，填入 DEEPSEEK_API_KEY")

    doc.add_heading("13.2 数据初始化", level=2)
    add_code_block(doc, "# 3. 将 CSV 导入 SQLite")
    add_code_block(doc, "python scripts/init_db.py")
    add_code_block(doc, "")
    add_code_block(doc, "# 4. 构建 TF-IDF 向量索引")
    add_code_block(doc, "python scripts/build_index.py")

    doc.add_heading("13.3 启动服务", level=2)
    add_code_block(doc, "# 启动 FastAPI 后端 (端口 8000)")
    add_code_block(doc, "uvicorn api:app --reload --port 8000")
    add_code_block(doc, "# API 文档: http://localhost:8000/docs")
    add_code_block(doc, "")
    add_code_block(doc, "# 启动 Streamlit 前端 (默认端口 8501)")
    add_code_block(doc, "streamlit run app.py")

    doc.add_heading("13.4 测试与评测", level=2)
    add_code_block(doc, "# 运行单元测试和集成测试")
    add_code_block(doc, "pytest tests/ -v")
    add_code_block(doc, "")
    add_code_block(doc, "# 运行评测 (需要 DeepSeek API Key)")
    add_code_block(doc, "python scripts/evaluate.py")
    add_code_block(doc, "")
    add_code_block(doc, "# 生成技术报告 (Word)")
    add_code_block(doc, "python scripts/generate_report.py")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 14: 技能亮点
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("十四、项目价值与技能亮点总结", level=1)

    doc.add_paragraph(
        "招投标智能体平台是一个完整的、可投入使用的 AI Agent 系统，体现了以下核心技术能力："
    )

    doc.add_heading("14.1 AI/LLM 工程能力", level=2)
    add_bullet(doc, "从零手写 Tool Calling 循环 (非 LangChain/AutoGPT)，深入理解 Agent 底层机制")
    add_bullet(doc, "Multi-Agent 编排 (Supervisor + Specialists)，具备跨领域工具编排能力")
    add_bullet(doc, "Prompt Engineering: 设计分层次的中文系统提示词，引导 LLM 准确进行工具选择和结果整合")
    add_bullet(doc, "DeepSeek API 集成，掌握 OpenAI 兼容 API 的使用模式")

    doc.add_heading("14.2 RAG / 检索引擎", level=2)
    add_bullet(doc, "从零实现 TF-IDF + Cosine Similarity 语义检索引擎")
    add_bullet(doc, "针对中文优化的 char_wb 字符 n-gram 分词策略")
    add_bullet(doc, "接口设计预留神经网络 Embedding 替换能力")

    doc.add_heading("14.3 全栈工程能力", level=2)
    add_bullet(doc, "FastAPI RESTful API 设计 (Pydantic 验证 + CORS + 会话管理)")
    add_bullet(doc, "Streamlit 多 Agent 对话前端 (Session State 管理 + 缓存策略)")
    add_bullet(doc, "SQLite 数据库设计 (表结构 + 索引策略 + 自动化导入管道)")
    add_bullet(doc, "Pandas 数据处理管道 (编码检测 + 中文列名标准化 + 类型推断)")

    doc.add_heading("14.4 工程质量", level=2)
    add_bullet(doc, "完整的测试体系: 24 个 pytest 用例 (8 单元测试 Mock + 16 集成测试)")
    add_bullet(doc, "50 条标注数据的自动化评测框架，支持 4 维度指标")
    add_bullet(doc, "指数退避重试、优雅降级、幂等设计等生产级工程实践")
    add_bullet(doc, "自动化技术文档生成 (python-docx Word 报告)")

    doc.add_heading("14.5 代码设计能力", level=2)
    add_bullet(doc, "清晰的分层架构: 展示层 → API 层 → Agent 编排层 → 工具层 → 数据层")
    add_bullet(doc, "工厂模式 (AgentRegistry)、单例模式 (VectorStore)、策略模式 (多编码检测)")
    add_bullet(doc, "接口预留设计，保证未来可扩展性")
    add_bullet(doc, "约 2000 行 Python 代码，单文件不超过 315 行，模块职责清晰")

    # ═══════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════
    doc.save(str(OUT_PATH))
    print(f"Report saved: {OUT_PATH}")
    print(f"Size: {os.path.getsize(OUT_PATH) / 1024:.0f} KB")


if __name__ == "__main__":
    build_document()
